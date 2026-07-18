from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document as DocxDocument
from pypdf import PdfWriter

from enterprise_doc_core.documents.ingestion import (
    DocumentParseViolation,
    HashEmbeddingProvider,
    ParsedSection,
    chunk_sections,
    parse_document_bytes,
)


def test_txt_parser_preserves_headings_and_offsets() -> None:
    sections = parse_document_bytes(b"# Terms\nPayment is due in 30 days.\n", extension=".txt")

    assert sections[0].heading == "Terms"
    assert sections[0].text == "Payment is due in 30 days."
    assert sections[0].start_offset == 8
    assert sections[0].end_offset > sections[0].start_offset


def test_docx_parser_extracts_heading_and_paragraph() -> None:
    document = DocxDocument()
    document.add_heading("Delivery", level=1)
    document.add_paragraph("Acceptance requires a written sign-off.")
    output = BytesIO()
    document.save(output)

    sections = parse_document_bytes(output.getvalue(), extension=".docx")

    assert [section.heading for section in sections] == ["Delivery"]
    assert "written sign-off" in sections[0].text


def test_pdf_parser_returns_page_metadata() -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    output = BytesIO()
    writer.write(output)

    sections = parse_document_bytes(output.getvalue(), extension=".pdf")

    assert sections[0].page_number == 1


def test_chunker_is_deterministic_and_bounded() -> None:
    sections = (
        ParsedSection(page_number=1, heading="A", text="abcdefghij", start_offset=0, end_offset=10),
    )

    first = chunk_sections(sections, max_chars=6, overlap_chars=2)
    second = chunk_sections(sections, max_chars=6, overlap_chars=2)

    assert first == second
    assert [chunk.chunk_index for chunk in first] == [0, 1, 2]
    assert all(len(chunk.text) <= 6 for chunk in first)


async def test_hash_embedding_provider_is_normalized_and_repeatable() -> None:
    provider = HashEmbeddingProvider(dimension=8)

    first = await provider.embed(("payment term", "payment term"))
    second = await provider.embed(("payment term", "payment term"))

    assert first == second
    assert len(first[0]) == 8
    assert abs(sum(value * value for value in first[0]) - 1.0) < 1e-6


@pytest.mark.parametrize(
    ("data", "extension", "code"),
    [
        (b"\xff", ".txt", "text_decode_failed"),
        (b"not a pdf", ".pdf", "pdf_parse_failed"),
        (b"not a docx", ".docx", "docx_parse_failed"),
        (b"content", ".rtf", "document_type_unsupported"),
    ],
)
def test_parser_reports_stable_failures(data: bytes, extension: str, code: str) -> None:
    with pytest.raises(DocumentParseViolation) as caught:
        parse_document_bytes(data, extension=extension)

    assert caught.value.code == code
    assert "content" not in caught.value.message


def test_empty_text_document_produces_no_sections_or_chunks() -> None:
    sections = parse_document_bytes(b"\n\n", extension="txt")

    assert sections == ()
    assert chunk_sections(sections, max_chars=10, overlap_chars=2) == ()


def test_chunker_preserves_metadata_offsets_and_hash() -> None:
    chunks = chunk_sections(
        (
            ParsedSection(
                page_number=3, heading="Payment", text="abcdef", start_offset=10, end_offset=16
            ),
        ),
        max_chars=4,
        overlap_chars=1,
    )

    assert chunks[0].page_number == 3
    assert chunks[0].heading == "Payment"
    assert chunks[0].start_offset == 10
    assert chunks[0].end_offset == 14
    assert len(chunks[0].content_sha256) == 64


@pytest.mark.parametrize(
    ("max_chars", "overlap_chars"),
    [(0, 0), (-1, 0), (4, -1), (4, 4), (4, 5)],
)
def test_chunker_rejects_invalid_bounds(max_chars: int, overlap_chars: int) -> None:
    with pytest.raises(ValueError):
        chunk_sections((), max_chars=max_chars, overlap_chars=overlap_chars)


def test_hash_embedding_provider_rejects_invalid_dimension() -> None:
    with pytest.raises(ValueError):
        HashEmbeddingProvider(dimension=0)
